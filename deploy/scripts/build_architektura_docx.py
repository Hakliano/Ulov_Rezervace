#!/usr/bin/env python3
"""Vygeneruje ARCHITEKTURA.docx z obsahu dokumentace."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "ARCHITEKTURA.docx"

BRAND = RGBColor(0x1A, 0x36, 0x5D)
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
MUTED = RGBColor(0x4A, 0x55, 0x68)
CODE_BG = "F1F5F9"


def set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(0)
        set_cell_shading(hdr_cells[i], "1A365D")

    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.strip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    # light background via shading on paragraph (approximation: indent + monospace)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_BG)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        for r in p.runs:
            r.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = BRAND if level == 1 else ACCENT
        r.font.name = "Calibri"


def add_body(doc: Document, text: str, bold_parts: list[str] | None = None) -> None:
    p = doc.add_paragraph()
    if not bold_parts:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        return
    # simple **bold** split
    rest = text
    for bp in bold_parts:
        if bp in rest:
            before, _, after = rest.partition(bp)
            if before:
                r = p.add_run(before)
                r.font.size = Pt(11)
            rb = p.add_run(bp)
            rb.bold = True
            rb.font.size = Pt(11)
            rest = after
    if rest:
        r = p.add_run(rest)
        r.font.size = Pt(11)


def build() -> Path:
    doc = Document()

    # stránka
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # titulní blok
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("ULOV KLIENTY")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = BRAND
    tr.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Architektura a provoz systému")
    sr.font.size = Pt(16)
    sr.font.color.rgb = ACCENT
    sr.font.name = "Calibri"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Stručný technický přehled · červenec 2026")
    mr.font.size = Pt(11)
    mr.font.color.rgb = MUTED
    doc.add_paragraph()

    add_body(
        doc,
        "Produkt: web na míru + online rezervace pro salony, provozovny, řemesla a zdraví. "
        "Repo: Hakliano/Ulov_Rezervace.",
    )
    add_body(
        doc,
        "Detailnější materiály: TECHNICKE_NASAZENI.md, NASAZENI_PRODUKCE.md, deploy/DEPLOY_PIPELINE.md.",
    )

    # 1
    add_heading(doc, "1. Co systém dělá", 1)
    add_table(
        doc,
        ["Vrstva", "Funkce"],
        [
            ["Prezentační web (presentace/)", "Marketing hub, vertikály, poptávkový formulář"],
            ["Web klienta (salonN/, vertikální dema)", "Úvod, galerie, personál, ceník, novinky, kontakt"],
            ["Rezervace (rezervace.html + JS)", "Výběr služby, personálu, termínu; e-mail; storno tokenem"],
            ["Administrace", "Obsah webu, kalendář, personál, NO-show, QR platby, statistiky"],
            ["Backend API (backend/)", "Multi-tenant logika, e-maily, uploady, GDPR životní cyklus"],
        ],
    )
    add_body(doc, "Jeden backend obsluhuje více salonů (tenantů). Data jsou izolovaná přes salon_id.")

    # 2
    add_heading(doc, "2. Architektura (přehled)", 1)
    add_code_block(
        doc,
        """┌─────────────────────────────────────────────────────────────────┐
│  Prohlížeč                                                      │
│  presentace/  salon1–8/  zdravi-*/ remesla-*/ provoz-*/         │
│  (HTML + CSS + vanilla JS, SALON_ID v app.js / rezervace.js)    │
└────────────────────────────┬────────────────────────────────────┘
                             │ HTTPS REST  →  api.ulovklienty.cz
                             ▼
┌─────────────────────────────────────────────────────────────────┐
│  Hetzner VPS — Docker Compose                                   │
│  nginx (TLS, statika) → Gunicorn → Django 5 + DRF              │
│  PostgreSQL 16 │ Redis 7 │ Celery worker │ cron                 │
└──────┬──────────────────┬──────────────────┬────────────────────┘
       ▼                  ▼                  ▼
  PostgreSQL          Bunny.net CDN       SMTP per salon""",
    )
    add_body(
        doc,
        "Staging běží na stejném serveru jako LIVE — oddělená DB (ulov_staging), "
        "statika v www-staging/, API kontejner ulov-staging-api.",
    )

    # 3
    add_heading(doc, "3. Technologie", 1)
    add_table(
        doc,
        ["Oblast", "Stack"],
        [
            ["Frontend", "Statické HTML/CSS/JS, responzivní layout"],
            ["API", "Django 5, Django REST Framework, CORS"],
            ["DB", "PostgreSQL 16 (produkce/staging), SQLite jen lokální dev"],
            ["Fronta", "Redis 7 + Celery (e-maily)"],
            ["Web server", "Nginx + Let's Encrypt (Certbot)"],
            ["Kontejnery", "Docker Compose"],
            ["Média", "Bunny.net Storage + CDN"],
            ["E-maily", "SMTP per salon; staging má EMAIL_OVERRIDE_TO"],
            ["Platby", "QR kódy SPAYD (české banky)"],
        ],
    )

    # 4
    add_heading(doc, "4. Struktura repozitáře", 1)
    add_table(
        doc,
        ["Složka", "Účel"],
        [
            ["backend/", "Django API — salons, rezervace, partner_admin"],
            ["salon1/ … salon8/", "Beauty dema (web + rezervace + admin)"],
            ["zdravi-*, remesla-*, provoz-*", "Vertikální dema"],
            ["presentace/", "Hub + landingy vertikál"],
            ["shared/", "Sdílené assety"],
            ["deploy/", "Deploy skripty, nginx, zálohy"],
            ["www/", "Výstup na serveru — sync při deployi (není zdroj pravdy)"],
        ],
    )
    add_body(doc, "Zdroj pravdy pro weby je git. Složka www/ na Hetzneru je jen nasazená kopie.")

    # 5
    add_heading(doc, "5. Napojení a URL", 1)
    add_heading(doc, "LIVE", 2)
    add_table(
        doc,
        ["Služba", "URL"],
        [
            ["Hub / presentace", "https://www.ulovklienty.cz/"],
            ["Vertikály", "/beauty/, /zdravi/, /remesla/, /provozovny/"],
            ["Vertikální dema", "/zdravi-fyzio/, /provoz-pujcovna/, …"],
            ["Beauty dema", "demo1.ulovklienty.cz … demo8 → salon1 … salon8"],
            ["API", "https://api.ulovklienty.cz/api/"],
            ["Health", "https://api.ulovklienty.cz/health/"],
        ],
    )
    add_heading(doc, "Staging", 2)
    add_table(
        doc,
        ["Služba", "URL"],
        [
            ["Hub", "https://www.staging.ulovklienty.cz/"],
            ["Dema", "https://www.staging.ulovklienty.cz/salon1/ …"],
            ["API", "https://api-staging.ulovklienty.cz/"],
        ],
    )
    add_heading(doc, "Externí služby", 2)
    add_bullets(
        doc,
        [
            "GitHub — verzování, větve dev / main",
            "Bunny.net — CDN pro loga, hero, galerie, upload z administrace",
            "SMTP — připomínky, potvrzení, recenze",
            "DNS — *.ulovklienty.cz → Hetzner 49.13.23.65",
        ],
    )
    add_heading(doc, "API konvence", 2)
    add_bullets(
        doc,
        [
            "Base: /api/salon/<id>/…",
            "Veřejné: čtení webu, volné termíny, vytvoření rezervace",
            "Admin zápis: hlavička X-Admin-Password",
            "Poptávka z presentace: POST /api/poptavka/",
        ],
    )

    # 6
    add_heading(doc, "6. Hlavní funkce backendu", 1)
    add_bullets(
        doc,
        [
            "salons — ceník, novinky, galerie, personál, otevírací doba, upload na Bunny",
            "rezervace — kalendář, zaměstnanci, rezervace, storno, NO-show, QR platby",
            "partner_admin — nastavení partnerů (domény, upozornění)",
            "Cron / Celery — připomínky, recenze, anonymizace a mazání dle GDPR",
        ],
    )

    # 7
    add_heading(doc, "7. Multi-tenant model", 1)
    add_bullets(
        doc,
        [
            "Jedna databáze, více salonů (Salon + salon_id u všech dat)",
            "Každý salon: vlastní branding, SMTP, admin heslo, Bunny prefix",
            "Zákaznická data a NO-show se mezi salony nesdílí",
            "Frontend má pevně SALON_ID v app.js / rezervace.js",
        ],
    )

    # 8
    add_heading(doc, "8. Deploy pipeline", 1)
    add_code_block(
        doc,
        "LOCAL (Cursor) → GitHub DEV → Staging → (schválení) → GitHub MAIN → LIVE",
    )
    add_table(
        doc,
        ["Krok", "Akce", "Kdo"],
        [
            ["①", "Commit + git push origin dev", "Po lokální kontrole"],
            ["②", "bash deploy/deploy-staging.sh origin/dev", "Na Hetzneru"],
            ["—", "Test na staging URL", "Ty"],
            ["③", "git push origin dev:main", "Po schválení stagingu"],
            ["④", "bash deploy/deploy-live.sh origin/main", "Jen po výslovném „na LIVE“"],
        ],
    )
    add_bullets(
        doc,
        [
            "Staging jen z dev, LIVE jen z main",
            "Žádný přímý scp/rsync z PC do www/",
            "Před LIVE: záloha (deploy/backup.sh), checklist (pre-deploy-check.sh)",
            "Havárie: deploy/rollback-live.sh",
        ],
    )

    # 9
    add_heading(doc, "9. Běžné použití (vývojář)", 1)
    add_heading(doc, "Lokální práce", 2)
    add_code_block(
        doc,
        """git checkout dev
git pull origin dev
# úpravy…
git add …
git commit -m "…"
git push origin dev""",
    )
    add_heading(doc, "Staging", 2)
    add_code_block(
        doc,
        """ssh root@49.13.23.65
cd /opt/ulov
bash deploy/deploy-staging.sh origin/dev""",
    )
    add_heading(doc, "LIVE (po schválení)", 2)
    add_code_block(
        doc,
        """git push origin dev:main
ssh root@49.13.23.65 'cd /opt/ulov && bash deploy/deploy-live.sh origin/main'""",
    )
    add_body(doc, "Po deployi API občas nginx vrátí 502 → docker compose exec -T nginx nginx -s reload.")
    add_body(doc, "MacBook na cestách: deploy/MACBOOK_CURSOR_SETUP.md.")

    # 10
    add_heading(doc, "10. Provoz na serveru", 1)
    add_table(
        doc,
        ["Cesta / služba", "Popis"],
        [
            ["/opt/ulov", "LIVE aplikace + git checkout"],
            ["docker compose", "db, api, worker, cron, nginx, certbot"],
            ["www/", "Nasazená statika"],
            ["www-staging/", "Staging statika"],
            ["backups/", "DB + config + www zálohy (14 dní)"],
        ],
    )

    # 11
    add_heading(doc, "11. Související dokumentace", 1)
    add_table(
        doc,
        ["Dokument", "Obsah"],
        [
            ["PREHLED_PRO_SALES_A_MARKETING.md", "Produkt pro obchod"],
            ["TECHNICKE_NASAZENI.md", "API, datový model"],
            ["NASAZENI_PRODUKCE.md", "Runbook nasazení"],
            ["deploy/DEPLOY_PIPELINE.md", "Staging, rollback"],
            ["deploy/DEPLOY_SAFETY.md", "Bezpečný sync"],
            ["deploy/MACBOOK_CURSOR_SETUP.md", "Setup MacBooku"],
        ],
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"OK: {path}")
